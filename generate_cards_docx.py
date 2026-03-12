"""
generate_cards_docx.py
Generates tent cards as Word (.docx) files using python-docx.
Each row in the Excel sheet becomes one tent card (page).
No circular imports — all helpers are defined locally.
"""

import os
import re
import datetime
import hashlib
import tempfile

import pandas as pd
import qrcode
from PIL import Image as PILImage

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# HELPERS — data cleaning (no import from generate_cards)
# ─────────────────────────────────────────────

def _clean(value):
    if pd.isna(value):
        return None
    value = str(value).strip()
    if value == "" or value.lower() == "nan":
        return None
    return value


def _clean_id(value):
    if pd.isna(value):
        return None
    value = str(value).strip()
    if value == "" or value.lower() == "nan":
        return None
    return value


_qr_cache = {}

def _create_qr(data):
    if data in _qr_cache:
        return _qr_cache[data]
    h    = hashlib.md5(str(data).encode()).hexdigest()[:12]
    path = os.path.join(tempfile.gettempdir(), f"qr_{h}.png")
    qrcode.make(str(data)).save(path)
    _qr_cache[data] = path
    return path


def _detect_patients(columns):
    nums = []
    for c in columns:
        m = re.match(r"patient (\d+) mrn", c.lower())
        if m:
            nums.append(int(m.group(1)))
    return sorted(set(nums))


def _get_location_fields(row, patient_number):
    suffix = "" if patient_number == 1 else str(patient_number)
    result = []
    for field in ["Site", "Building", "NurseStation", "Room", "Bed"]:
        val = _clean(row.get(f"{field}{suffix}"))
        if val:
            result.append((field.replace("NurseStation", "Ward"), val))
    return result


# ─────────────────────────────────────────────
# HELPERS — docx formatting
# ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_borders(cell, color='aaaaaa', sz=4):
    tc        = cell._tc
    tcPr      = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_table_borders(table, color='aaaaaa', sz=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_row_height(row, height_inches):
    tr   = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'),   str(int(height_inches * 1440)))
    trH.set(qn('w:hRule'), 'atLeast')
    trPr.append(trH)


def _para_bottom_border(para, color='aaaaaa', sz=4):
    """Add a bottom border to a paragraph — acts as a divider line."""""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(sz))
    bot.set(qn('w:color'), color)
    bot.set(qn('w:space'), '1')
    pBdr.append(bot)
    pPr.append(pBdr)


def _cant_split_table(table):
    """Prevent any row in the table from splitting across pages."""
    for row in table.rows:
        tr   = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        cs = OxmlElement('w:cantSplit')
        cs.set(qn('w:val'), '1')
        trPr.append(cs)


def _keep_with_next(para):
    """Keep this paragraph with the next element (heading stays with table)."""
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    kwn.set(qn('w:val'), '1')
    pPr.append(kwn)


def _bold_value(para, label, value, font_size=9):
    r1 = para.add_run(label)
    r1.bold       = True
    r1.font.size  = Pt(font_size)
    r2 = para.add_run(str(value))
    r2.bold       = False
    r2.font.size  = Pt(font_size)


def _add_qr(cell, qr_path, size_inches=0.85):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.add_run().add_picture(qr_path, width=Inches(size_inches))


def _spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def build_docx(df, sheet, logo_path=None, default_password=None,
               patient_labels=None, wristband_label=None):

    os.makedirs("output", exist_ok=True)
    ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"output/tent_cards_{ts}.docx"

    doc = Document()

    # Page: A4, 0.75in margins
    sec               = doc.sections[0]
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(0.75)
    sec.right_margin  = Inches(0.75)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)

    # Remove default spacing
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

    CW     = 6.77   # content width in inches
    HALF   = CW / 2
    cols   = list(df.columns)
    first  = True

    for _, row in df.iterrows():

        if not first:
            doc.add_page_break()
        first = False

        title = _clean(row.get("CLASS")) or sheet

        # ── 1. HEADER ────────────────────────────────────────
        h_tbl = doc.add_table(rows=1, cols=2)
        h_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Outer border only — no vertical divider between title and logo
        tbl_   = h_tbl._tbl
        tblPr_ = tbl_.find(qn('w:tblPr'))
        if tblPr_ is None:
            tblPr_ = OxmlElement('w:tblPr')
            tbl_.insert(0, tblPr_)
        tblB = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '6')
            el.set(qn('w:color'), '999999')
            tblB.append(el)
        for side in ['insideH', 'insideV']:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tblB.append(el)
        tblPr_.append(tblB)
        h_tbl.columns[0].width = Inches(CW - 1.8)
        h_tbl.columns[1].width = Inches(1.8)

        tc = h_tbl.cell(0, 0)
        _set_cell_bg(tc, 'e8e8e8')
        _set_cell_margins(tc, top=120, bottom=120, left=150, right=80)
        tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tp = tc.paragraphs[0]
        r  = tp.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

        lc = h_tbl.cell(0, 1)
        _set_cell_bg(lc, 'e8e8e8')
        _set_cell_margins(lc, top=60, bottom=60, left=60, right=80)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if logo_path and os.path.exists(logo_path):
            with PILImage.open(logo_path) as pil:
                ow, oh = pil.size
            ratio  = min((1.6 * 96) / ow, (0.45 * 96) / oh, 1.0)
            lp.add_run().add_picture(logo_path, width=Inches((ow * ratio) / 96))

        # ── 2. LOGIN BOXES ───────────────────────────────────
        login_data = []
        for i in range(1, 4):
            u   = _clean(row.get(f"Log On {i}"))
            pwd = default_password if default_password else _clean(row.get(f"Password {i}"))
            if u:
                login_data.append((u, pwd or ''))

        if login_data:
            _spacer(doc, 6)
            n      = len(login_data)
            GAP    = 0.08
            box_w  = (CW - GAP * (n - 1)) / n
            l_tbl  = doc.add_table(rows=1, cols=n)
            _remove_table_borders(l_tbl)
            l_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            for idx, (u, p) in enumerate(login_data):
                cell = l_tbl.cell(0, idx)
                cell.width = Inches(box_w)
                _set_cell_bg(cell, 'fafafa')
                _set_cell_borders(cell, color='aaaaaa', sz=4)
                _set_cell_margins(cell, top=60, bottom=60, left=110, right=60)

                p1 = cell.paragraphs[0]
                p1.paragraph_format.space_before = Pt(0)
                p1.paragraph_format.space_after  = Pt(0)
                _bold_value(p1, "User Name: ", u)

                p2 = cell.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after  = Pt(0)
                _bold_value(p2, "Password: ", p)

        # ── 3. PATIENT BLOCKS ────────────────────────────────
        patients = _detect_patients(cols)

        for n in patients:
            mrn = _clean_id(row.get(f"Patient {n} MRN"))
            if not mrn:
                continue
            fin   = _clean_id(row.get(f"Patient {n} FIN"))
            last  = _clean(row.get(f"Patient {n} Last Name"))
            first_name = _clean(row.get(f"Patient {n} First Name"))
            name  = f"{last or ''}, {first_name or ''}".strip().strip(',')

            # Patient label priority:
            # 1) Excel column PatientNLabel  2) UI override  3) default PATIENT
            excel_label = _clean(row.get(f"Patient{n}Label"))
            if excel_label:
                pat_label = excel_label.strip().upper()
            elif patient_labels and len(patient_labels) >= n and patient_labels[n-1].strip():
                pat_label = patient_labels[n-1].strip().upper()
            else:
                pat_label = "PATIENT"

            loc = _get_location_fields(row, n)
            _spacer(doc, 4)

            p_tbl = doc.add_table(rows=1, cols=2)
            _set_table_borders(p_tbl, color='aaaaaa', sz=4)
            p_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            left = p_tbl.cell(0, 0)
            left.width = Inches(HALF)
            _set_cell_margins(left, top=100, bottom=100, left=150, right=80)
            left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para_data in [(f"{pat_label}: ", name.upper()), ("MRN: ", mrn), ("FIN: ", fin or '')]:
                p_ = left.paragraphs[0] if para_data[0].startswith(pat_label) else left.add_paragraph()
                p_.paragraph_format.space_before = Pt(0)
                p_.paragraph_format.space_after  = Pt(2)
                _bold_value(p_, para_data[0], para_data[1])

            right = p_tbl.cell(0, 1)
            right.width = Inches(HALF)
            _set_cell_margins(right, top=100, bottom=100, left=150, right=80)
            right.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            lh = right.paragraphs[0]
            lh.paragraph_format.space_before = Pt(0)
            lh.paragraph_format.space_after  = Pt(2)
            r2 = lh.add_run("LOCATION")
            r2.bold = True
            r2.font.size = Pt(10)

            for i in range(0, len(loc), 2):
                pair  = loc[i:i+2]
                lp2   = right.add_paragraph()
                lp2.paragraph_format.space_before = Pt(0)
                lp2.paragraph_format.space_after  = Pt(2)
                for j, (lbl, val) in enumerate(pair):
                    if j > 0:
                        sep = lp2.add_run(",  ")
                        sep.font.size = Pt(9)
                    _bold_value(lp2, f"{lbl}: ", val)

        # ── 4. WRISTBAND QR ──────────────────────────────────
        wrist = []
        for n in patients:
            visit = _clean_id(row.get(f"Visit ID{n}"))
            if visit:
                wb = (wristband_label or "Patient Wristband").strip()
                wrist.append((f"{wb} {n}", _create_qr(visit)))

        if wrist:
            _spacer(doc, 6)
            ncols  = len(wrist)
            col_w  = CW / ncols
            wb_tbl = doc.add_table(rows=2, cols=ncols)
            _set_table_borders(wb_tbl, color='aaaaaa', sz=4)
            wb_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            for idx, (label, qr_path) in enumerate(wrist):
                hc = wb_tbl.cell(0, idx)
                hc.width = Inches(col_w)
                _set_cell_bg(hc, 'f5f5f5')
                _set_cell_margins(hc, top=80, bottom=80, left=80, right=80)
                hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                hp = hc.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hr = hp.add_run(label)
                hr.bold = True
                hr.font.size = Pt(10)

                qc = wb_tbl.cell(1, idx)
                qc.width = Inches(col_w)
                _set_cell_margins(qc, top=60, bottom=60, left=60, right=60)
                _add_qr(qc, qr_path, size_inches=1.0)

        # ── 5. MEDICATIONS ───────────────────────────────────
        meds = []
        for i in range(1, 15):
            mn  = _clean(row.get(f"Product{i}"))
            ndc = _clean_id(row.get(f"NDC{i}"))
            if mn and ndc:
                meds.append((mn, _create_qr(ndc)))

        if meds:
            _spacer(doc, 6)
            mh = doc.add_paragraph("Medications")
            mh.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mh.paragraph_format.space_after = Pt(4)
            mh.runs[0].bold = True
            mh.runs[0].font.size = Pt(11)
            _keep_with_next(mh)  # heading never strands without first table

            for i in range(0, len(meds), 2):
                pair = meds[i:i+2]
                while len(pair) < 2:
                    pair.append(None)
                # Single row — name + QR both inside each cell so they NEVER split
                mt = doc.add_table(rows=1, cols=2)
                _set_table_borders(mt, color='aaaaaa', sz=4)
                mt.alignment = WD_TABLE_ALIGNMENT.LEFT
                _cant_split_table(mt)

                for ci, med in enumerate(pair):
                    cell = mt.cell(0, ci)
                    cell.width = Inches(HALF)
                    _set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                    # Name paragraph — with bottom border as divider line
                    np2 = cell.paragraphs[0]
                    np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    np2.paragraph_format.space_before = Pt(4)
                    np2.paragraph_format.space_after  = Pt(4)
                    if med:
                        np2.add_run(med[0]).font.size = Pt(9)
                    _para_bottom_border(np2)

                    # QR paragraph — inside same cell, below name
                    if med:
                        qp = cell.add_paragraph()
                        qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        qp.paragraph_format.space_before = Pt(0)
                        qp.paragraph_format.space_after  = Pt(4)
                        qp.add_run().add_picture(med[1], width=Inches(0.85))

        # ── 6. SPECIMENS ─────────────────────────────────────
        specs = []
        for i in range(1, 10):
            sp  = _clean(row.get(f"Specimen{i}"))
            bc  = _clean_id(row.get(f"Barcode{i}"))
            if sp and bc:
                specs.append((sp, _create_qr(bc)))

        if specs:
            _spacer(doc, 6)
            sh = doc.add_paragraph("Specimens")
            sh.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sh.paragraph_format.space_after = Pt(4)
            sh.runs[0].bold = True
            sh.runs[0].font.size = Pt(11)
            _keep_with_next(sh)  # heading never strands without first table

            for i in range(0, len(specs), 2):
                pair = specs[i:i+2]
                while len(pair) < 2:
                    pair.append(None)
                # Single row — name + QR both inside each cell so they NEVER split
                st2 = doc.add_table(rows=1, cols=2)
                _set_table_borders(st2, color='aaaaaa', sz=4)
                st2.alignment = WD_TABLE_ALIGNMENT.LEFT
                _cant_split_table(st2)

                for ci, spec in enumerate(pair):
                    cell = st2.cell(0, ci)
                    cell.width = Inches(HALF)
                    _set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                    # Name paragraph — with bottom border as divider line
                    np3 = cell.paragraphs[0]
                    np3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    np3.paragraph_format.space_before = Pt(4)
                    np3.paragraph_format.space_after  = Pt(4)
                    if spec:
                        np3.add_run(spec[0]).font.size = Pt(9)
                    _para_bottom_border(np3)

                    # QR paragraph — inside same cell, below name
                    if spec:
                        qp = cell.add_paragraph()
                        qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        qp.paragraph_format.space_before = Pt(0)
                        qp.paragraph_format.space_after  = Pt(4)
                        qp.add_run().add_picture(spec[1], width=Inches(0.85))

    doc.save(filename)
    print("DOCX Generated:", filename)
    return filename
